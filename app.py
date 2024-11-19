import streamlit as st
import os
import io
import fitz
from google.cloud import vision
from google.cloud.vision_v1 import types
from openai import OpenAI
import tempfile
import json
import atexit
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import base64
import time
from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_ocr_locale import ExportOCRLocale
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult 




############# CHECK TO HANDLE DEPLOYED AND LOCAL ENVIRONMENT #############

# Check if we're running on Streamlit Cloud

if hasattr(st, "secrets"):

    # Deployed environment setup

    # Google Cloud setup
    google_creds = dict(st.secrets["GOOGLE_APPLICATION_CREDENTIALS"])
    google_creds_file = tempfile.NamedTemporaryFile(mode='w+', delete=False)
    json.dump(google_creds, google_creds_file)
    google_creds_file.flush()
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = google_creds_file.name
    
    # Adobe API setup
    os.environ["ADOBE_CLIENT_ID"] = st.secrets["ADOBE_CLIENT_ID"]
    os.environ["ADOBE_CLIENT_SECRET"] = st.secrets["ADOBE_CLIENT_SECRET"]
    
    # Add cleanup function and register it
    def cleanup_temp_files():
        if os.path.exists(google_creds_file.name):
            os.unlink(google_creds_file.name)
    
    atexit.register(cleanup_temp_files)
    
    # Initialize OpenAI client with Streamlit secrets
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

else:
    # Local environment setup

    from dotenv import load_dotenv
    load_dotenv()
    
    # Google Cloud Vision API for OCR setup
    os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')

    # Adobe API for PDF to Word setup
    os.environ['ADOBE_CLIENT_ID'] = os.getenv('ADOBE_CLIENT_ID')
    os.environ['ADOBE_CLIENT_SECRET'] = os.getenv('ADOBE_CLIENT_SECRET')
    
    # Initialize OpenAI client for gpt calls
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))




######################## STREAMLIT APP #########################

# Simple sidebar navigation
st.sidebar.title("IA - AVOCATS")

# Add horizontal white line and spacing after title
st.sidebar.markdown("""
    <hr style="border: 0.5px solid grey;">
""", unsafe_allow_html=True)

# Add custom CSS to increase spacing between radio options
st.markdown("""
    <style>
        div.row-widget.stRadio > div {
            flex-direction: column;
            gap: 25px;  /* Increase this value to add more space */
        }
    </style>
""", unsafe_allow_html=True)

# Radio button selection
selected_function = st.sidebar.radio(
    label="RÉSUMÉR ET CONCLURE",
    options=["RÉSUMÉ DE PIÈCES JURIDIQUES AVEC BORDEAUX", "PDF À WORD", "RÉSUMÉ DE DOCUMENT PDF OU WORD"],
    label_visibility="collapsed"
)

############### GÉNÉRER UN RÉSUMÉ DE PIÈCES JURIDIQUES AVEC BORDEAUX ###############
if selected_function == "RÉSUMÉ DE PIÈCES JURIDIQUES AVEC BORDEAUX":

    ################ LEGAL SUMMARIES AND IMAGE DESCRIPTIONS PROMPT TEMPLATES ################   

    # Summarize each legal OCR extracted transcript
    prompt_template_summary = """
    Résume le texte entre les triples parenthèses en suivant ces directives :

    - Extrait les faits et informations importantes à mentionner
    - Résume et output en français.
    - Accorde les verbes au passé composé ou à l'imparfait.
    - Extrait et ajoute la date en format "JJ mois AAAA" 

    Met la totalité du texte en forme une seule fois en suivant cette structure :

    Le <date>, <brève explication des faits>.

    <Explication résumée des faits en utilisant les termes clés>.

    ((({})))"""

    # Generate the bordereau entries for each OCR transcript + image descriptions
    prompt_template_bordereau = """
    Tu reçois un transcript d'une pièce juridique entre les triples parenthèses.
    Tu vas générer une ligne de bordereau de pièce en suivant ces directives :

    - Extrait le titre de la pièce qui décrit le plus justement la pièce en utilisant la terminologie juridique.
    - Sois précis et concis dans le titre.
    - Écris le titre avec une majuscule au début et le reste en minuscules.
    - Ne mentionne pas le numéro du document dans le titre.

    Output en suivant cette structure :
    <Titre de la pièce>

    IMPORTANT :
    - Relis ton output et verifie les conditions suivantes :
    - Dans le cas précis où la pièce est une attestation de témoin, mentionne le genre (Monsieur ou Madame) et le nom de famille de l'individu seulement (tout en majuscules).

    ((({})))"""

    # Image description with gpt-4o 
    prompt_template_image = """

    Tu reçois une image de document juridique en rapport avec une affaire. 
    Décris cette image de document juridique en français de manière concise et factuelle.
    Retiens uniquement les éléments importants. 

    Commence ta description par "La pièce image montre" et reste bref.
    Ouput 2 à 3 phrases maximum.
    """

    # Generate a title from the image description
    prompt_template_image_title = """

    Génère un titre pour la description des images entres les triples parenthèses.
    Le titre doit être court et significatif.
    Output en une seule phrase.
    Output en français.
    Commence par une majuscule et finis par un point.

    Exemple :
    <Titre à générer>

    ((({})))
    """

    # Classify the page as TEXT, IMAGE, or SKIP
    prompt_template_classification = """
    Analyze this page and classify it as either:
    1. "TEXT" - if it contains meaningful text content that should be processed with OCR
    2. "IMAGE" - if it's primarily an image, photo, ID, or visual document that needs description
    3. "SKIP" - if it's a blank or nearly blank page with no meaningful content

    It is crucial and extremely important that you output ONLY with either "TEXT", "IMAGE", or "SKIP"
    """

    # Function to process uploaded files and generate summaries and bordereau
    def process_uploaded_files(uploaded_files):
        """Process PDFs and generate summaries and bordereau."""
        client = vision.ImageAnnotatorClient()
        all_summaries = []
        bordereau_entries = []
        
        # Create progress bar and status message
        progress_bar = st.progress(0)
        status_message = st.empty()
        
        total_files = len(uploaded_files)
        
        for index, pdf_file in enumerate(uploaded_files):
            # Update status for current file
            current_progress = int((index / total_files) * 70)  # 70% of progress for PDF processing
            status_message.write(f"L'IA traite les PDFs... ({index + 1}/{total_files} fichiers)")
            progress_bar.progress(current_progress)
            
            # Extract piece number from filename
            piece_num = re.search(r'\D*(\d+)', pdf_file.name)
            piece_num = piece_num.group(1) if piece_num else "X"
            print(f"Processing file: {pdf_file.name} as Pièce nº{piece_num}")
            
            # Process PDF
            pdf_content = pdf_file.read()
            pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
            
            transcript = []
            image_descriptions = []
            
            # Process each page
            for page in pdf_document:
                # Convert page to image
                zoom = 300 / 72  # 300 DPI
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                img_bytes = pix.tobytes()
                
                # Get text using Google Vision OCR
                image = types.Image(content=img_bytes)
                response = client.document_text_detection(image=image)
                page_text = response.full_text_annotation.text if response.full_text_annotation else ""
                
                # Process based on content length
                if len(page_text) > 700:
                    transcript.append(page_text)
                else:
                    # Classify page with GPT
                    try:
                        base64_image = base64.b64encode(img_bytes).decode('utf-8')
                        classification = process_with_gpt(
                            prompt=prompt_template_classification,
                            image_base64=base64_image,
                            is_classification=True
                        )
                        
                        if "TEXT" in classification:
                            transcript.append(page_text)
                        elif "IMAGE" in classification:
                            # Get image description
                            description = process_with_gpt(
                                prompt=prompt_template_image,
                                image_base64=base64_image,
                                is_image_description=True
                            )
                            if description:
                                image_descriptions.append(description)
                    except Exception as e:
                        print(f"Error processing page: {e}")
            
            # Create summary for this PDF
            full_transcript = '\n\n'.join(transcript)
            if full_transcript:
                # Summarize transcript
                summary = process_with_gpt(
                    prompt=prompt_template_summary.format(full_transcript)
                )
                if summary:
                    # Add image descriptions and piece number
                    if image_descriptions:
                        desc_text = "\n\n".join(image_descriptions)
                        summary = f"{summary}{desc_text} (Pièce nº{piece_num})"
                    else:
                        summary = f"{summary} (Pièce nº{piece_num})"
            else:
                # Images-only piece
                if image_descriptions:
                    images_text = "\n\n".join(image_descriptions)
                    title = process_with_gpt(
                        prompt=prompt_template_image_title.format(images_text)
                    )
                    title = title if title else "Images"
                    summary = f"Le JJ mois AAAA, {title}\n\n{images_text} (Pièce nº{piece_num})"
                else:
                    summary = f"Pièce vide (Pièce nº{piece_num})"
            
            all_summaries.append(summary)
            
            # NEW: Extract date from first line of summary
            first_line = summary.strip().split('\n')[0]
            date_match = re.match(r'Le (\d{1,2} \w+ \d{4})', first_line)
            extracted_date = date_match.group(1) if date_match else "JJ mois AAAA"
            
            # Generate bordereau entry
            combined_text = full_transcript
            if image_descriptions:
                combined_text += "\n\n".join(image_descriptions)
            
            bordereau_entry = process_with_gpt(
                prompt=prompt_template_bordereau.format(combined_text)
            )
            if bordereau_entry:
                # Format bordereau entry with piece number, title, and date
                bordereau_entry = f"{piece_num} - {bordereau_entry} - du {extracted_date}"
                bordereau_entries.append(bordereau_entry)
        
        # Start the chronological sorting progress bar
        status_message.write("Tri chronologique des résumés...")
        progress_bar.progress(70)  # 70% for chronological sorting
        time.sleep(2)  # Add small delay for visibility
        
        # Combine all results
        combined_summaries = "\n\n------\n\n".join(all_summaries)
        chronological_summary = sort_summaries_chronologically(combined_summaries)
        
        # Update status for chronological sorting
        status_message.write("Finalisation...")
        progress_bar.progress(85)  # 85% to finalize
        time.sleep(2)  # Add small delay for visibility
        
        # Create bordereau section
        bordereau_section = "BORDEREAU DE PIECES COMMUNIQUEES\n\n" + "\n".join(entry + "\n" for entry in bordereau_entries)

        # Final status update
        status_message.write("Fini!")
        progress_bar.progress(100)
        time.sleep(2)  # Add small delay for visibility
        
        # Clear progress indicators
        progress_bar.empty()
        status_message.empty()
        
        return {
            'original': f"{combined_summaries}\n\n{'='*50}\n\n{bordereau_section}",
            'chronological': f"{chronological_summary}\n\n{'='*50}\n\n{bordereau_section}"
        }

    # Function to handle all GPT API calls, with or without images
    def process_with_gpt(prompt, image_base64=None, is_classification=False, is_image_description=False):
        """Handle all GPT API calls, with or without images."""
        try:
            messages = [{
                "role": "user",
                "content": [{"type": "text", "text": prompt}]
            }]
            
            # Add image to message if provided
            if image_base64:
                messages[0]["content"].append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{image_base64}",
                        "detail": "high"
                    }
                })
            
            # Use gpt-4o for image-related tasks, gpt-4o-mini for text-only tasks
            model = "gpt-4o" if image_base64 else "gpt-4o-mini-2024-07-18"
            
            # Temperature 1 only for image descriptions, 0 for everything else
            temperature = 1 if is_image_description else 0
            
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature
            )
            
            result = response.choices[0].message.content.strip()
            
            # For classification, return uppercase result
            if is_classification:
                return result.upper()
            return result
            
        except Exception as e:
            print(f"Error in GPT processing: {e}")
            return None
    
    # Function to sort summaries based only on their initial paragraph dates
    def sort_summaries_chronologically(combined_summaries):
        """
        Sort summaries based only on their initial paragraph dates.
        """
        summaries = combined_summaries.split('------\n\n')
        dated_summaries = []
        undated_summaries = []
        
        for summary in summaries:
            summary = summary.strip()
            date = parse_initial_date_fr(summary)
            
            if date:
                dated_summaries.append((date, summary))
            else:
                undated_summaries.append(summary)
        
        dated_summaries.sort(key=lambda x: x[0])
        sorted_texts = [s[1] for s in dated_summaries] + undated_summaries
        return '\n\n------\n\n'.join(sorted_texts)

    # Function to parse ONLY the date at the start of a paragraph
    def parse_initial_date_fr(text):
        """
        Parse ONLY the date at the start of a paragraph.
        Returns datetime object or None if no valid date found.
        """
        fr_months = {
            'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4,
            'mai': 5, 'juin': 6, 'juillet': 7, 'août': 8,
            'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12
        }
        
        first_line = text.strip().split('\n')[0]
        pattern = r'^Le (\d{1,2}) (\w+) (\d{4})'
        match = re.match(pattern, first_line)
        
        if match:
            day, month_fr, year = match.groups()
            month_num = fr_months.get(month_fr.lower())
            if month_num:
                try:
                    return datetime(int(year), month_num, int(day))
                except ValueError:
                    return None
        return None

    # Function to create a Word document from a passed text
    def create_summary_word_document(summary_text):
        doc = Document()
        
        title = doc.add_paragraph("Résumé des pièces")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        doc.add_paragraph(summary_text)
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer


    st.title("RÉSUMÉ DE PIÈCES JURIDIQUES PDF ET GÉNÉRATION DE BORDEAUX")
    
    st.markdown("\n")
    st.markdown("#####  A QUOI SERT CET OUTIL ?")
    st.markdown("\n")
    st.markdown("#####  Gagnez en temps et compréhension dans la préparation de vos conclusions grâce à l'IA en :")
    st.markdown("######  - Résumant les pièces automatiquement (dans leur ordre original et chronologique).")
    st.markdown("######  - Décrivant les images des pièces.")
    st.markdown("######  - Générant un bordereau.")
    st.markdown("######  - Permettant un copier-coller facile dans vos conclusions après relecture.")


    st.markdown("####")

    st.markdown("#####  ÉTAPES À SUIVRE POUR UN RÉSULTAT OPTIMAL :")
    st.markdown("\n")    
    st.markdown("#####  1. Préparez et uploadez la totalité de vos pièces juridiques en PDF en suivant les directives ci-dessous :")
    st.markdown("- Une seule pièce juridique par PDF (chaque fichier peut contenir plusieurs pages).")
    st.markdown("- Chaque PDF doit être nommé avec son numéro de pièce. (Ex. : Pièce nº1.pdf,  Piece2.pdf,  3.pdf,  etc.)")
    st.markdown("- La meilleure qualité de document possible est requise pour un résultat optimal.")
    st.markdown("#####  2. L'IA générera le résumé (avec sa version chronologique) et le bordereau.")
    st.markdown("#####  3. Téléchargez le résultat en format Word ou .txt.")

    st.markdown("####")

    st.markdown("#####  IMPORTANT :")
    st.markdown("L'IA commet des erreurs de temps en temps. Une relecture est donc nécessaire.")
    st.markdown("Cet outil est là pour le gain de temps à l'étape de rédaction, et vous permettre de vous concentrer sur l'exactitude des résumés et conclusions.")
    st.markdown("En fonction de la qualité des PDFs et du type de contenu, certains résumés peuvent être incomplets ou contenir des erreurs.")
    st.markdown("L'IA s'efforce de suivre une structure adaptée aux conclusions, cependant, certaines pièces manqueront d'informations pour cette tâche.")    


    st.markdown("####")

    # Initialize session state
    if 'summaries' not in st.session_state:
        st.session_state.summaries = None

    st.markdown("#### Ajoutez toutes vos pièces juridiques en PDF ci-dessous")
    uploaded_files = st.file_uploader(
        label="Upload des pièces juridiques",
        type="pdf", 
        accept_multiple_files=True, 
        label_visibility="collapsed",
        key="pdf_uploader_summary"
    )

    if uploaded_files:
        if st.button("Générer le résumé"):
            with st.spinner("Nous générons votre résumé..."):
                # Sort files using natural number ordering
                sorted_files = sorted(uploaded_files, 
                    key=lambda f: (
                        float('inf') if not re.search(r'\D*(\d+)', f.name) 
                        else int(re.search(r'\D*(\d+)', f.name).group(1))
                    ))
                st.session_state.summaries = process_uploaded_files(sorted_files)

    if st.session_state.summaries:
        st.markdown("### ")
        st.markdown("### Résumé et Bordereau")
        
        # Add tabs for different versions
        tab1, tab2 = st.tabs(["Version originale", "Version chronologique"])
        
        with tab1:
            st.markdown(st.session_state.summaries['original'])
            
            # Create two columns for download buttons
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📥 Télécharger en .txt",
                    data=st.session_state.summaries['original'],
                    file_name="resume_et_bordereau.txt",
                    mime="text/plain",
                    key="download_original_txt"
                )
            
            with col2:
                word_buffer = create_summary_word_document(st.session_state.summaries['original'])
                st.download_button(
                    label="📥 Télécharger en Word",
                    data=word_buffer,
                    file_name="resume_et_bordereau.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_original_word"
                )
        
        with tab2:
            st.markdown(st.session_state.summaries['chronological'])
            
            # Create two columns for download buttons
            col1, col2 = st.columns(2)
            
            with col1:
                st.download_button(
                    label="📥 Télécharger en .txt",
                    data=st.session_state.summaries['chronological'],
                    file_name="resume_et_bordereau_chronologique.txt",
                    mime="text/plain",
                    key="download_chrono_txt"
                )
            
            with col2:
                word_buffer = create_summary_word_document(st.session_state.summaries['chronological'])
                st.download_button(
                    label="📥 Télécharger en Word",
                    data=word_buffer,
                    file_name="resume_et_bordereau_chronologique.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_chrono_word"
                )


######################### CONVERTISSEUR PDF VERS WORD #########################
elif selected_function == "PDF À WORD":

    def convert_pdf_to_word(uploaded_file):
        """
        Convert PDF to Word using Adobe PDF Services API with OCR support.
        """
        try:
            # Initialize credentials and service
            credentials = ServicePrincipalCredentials(
                client_id=os.environ["ADOBE_CLIENT_ID"],
                client_secret=os.environ["ADOBE_CLIENT_SECRET"]
            )
            pdf_services = PDFServices(credentials=credentials)

            # Upload PDF and set conversion parameters
            input_asset = pdf_services.upload(
                input_stream=uploaded_file.getvalue(), 
                mime_type=PDFServicesMediaType.PDF
            )
            export_params = ExportPDFParams(
                target_format=ExportPDFTargetFormat.DOCX,
                ocr_lang=ExportOCRLocale.FR_FR
            )

            # Convert PDF to Word
            export_job = ExportPDFJob(input_asset=input_asset, export_pdf_params=export_params)
            location = pdf_services.submit(export_job)
            result = pdf_services.get_job_result(location, ExportPDFResult)

            # Prepare for download
            word_buffer = io.BytesIO(
                pdf_services.get_content(
                    result.get_result().get_asset()
                ).get_input_stream()
            )
            word_buffer.seek(0)
            
            return word_buffer

        except Exception as e:
            st.error(f"Erreur lors de la conversion: {str(e)}")
            return None

    st.title("CONVERTISSEUR PDF VERS WORD")
 

    st.markdown("\n")
    st.markdown("#####  Uploadez un fichier PDF et obtenez sa conversion en document Word.")

    st.markdown("\n")
    st.markdown("####")
    st.markdown("#####  IMPORTANT")
    st.markdown("######  - Le PDF doit avoir une bonne qualité d'image pour une meilleure conversion.")
    st.markdown("######  - Plus le document contient d'éléments non textuels, moins fidèle sera la conversion.")
    st.markdown("######  - La conversion peut prendre quelques minutes selon le nombre de pages.")
    st.markdown("####")
    
    st.markdown("#### Ajoutez votre fichier PDF ci-dessous")
    uploaded_file = st.file_uploader(
        label="Upload du PDF",
        type="pdf", 
        accept_multiple_files=False, 
        label_visibility="collapsed",
        key="pdf_uploader_word"
    )
    
    if uploaded_file:
        if st.button("Convertir en Word"):
            with st.spinner("Conversion en cours... Cela peut prendre quelques minutes."):
                word_buffer = convert_pdf_to_word(uploaded_file)
                if word_buffer:
                    st.success("Conversion réussie!")
                    st.download_button(
                        label="📥 Télécharger le document Word",
                        data=word_buffer,
                        file_name=f"{os.path.splitext(uploaded_file.name)[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


######################### RÉSUMÉ SIMPLE DE DOCUMENT PDF OU WORD #########################
elif selected_function == "RÉSUMÉ DE DOCUMENT PDF OU WORD":

    # General summarization prompt template for single documents
    prompt_template_general = """

    Résume le texte entre les triples parenthèses:
    N'ajoute pas de titre ou de conclusion.
    N'ajoute pas "Résumé" ou "Summary" au début.
    N'ajoute pas "Le texte" au début.

    ((({})))
    """

    def create_single_document_summary(uploaded_file):
        """Process PDF or Word document and generate summary."""
        client = vision.ImageAnnotatorClient()
        all_chunks_summaries = []
        MAX_TOKENS = 1000  # Token limit for GPT
        CHARS_PER_TOKEN = 4  # Approximate chars per token
        MAX_CHUNK_SIZE = MAX_TOKENS * CHARS_PER_TOKEN
        
        # Create progress bar and status message
        progress_bar = st.progress(0)
        status_message = st.empty()
        
        try:
            # Extract text based on file type
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension in ['doc', 'docx']:
                # Process Word document
                status_message.write("L'IA extrait le texte...")
                doc = Document(uploaded_file)
                full_text = '\n\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
                
            elif file_extension == 'pdf':
                # Process PDF with OCR
                status_message.write("L'IA extrait le texte...")
                pdf_content = uploaded_file.read()
                pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
                
                full_text = []
                total_pages = len(pdf_document)
                
                for page_num, page in enumerate(pdf_document):
                    # Update progress
                    progress = int((page_num / total_pages) * 50)  # First 50% for OCR
                    progress_bar.progress(progress)
                    status_message.write(f"L'IA extrait le texte... (page {page_num + 1}/{total_pages})")
                    
                    # Convert page to image
                    zoom = 300 / 72  # 300 DPI
                    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                    img_bytes = pix.tobytes()
                    
                    # Get text using Google Vision OCR
                    image = types.Image(content=img_bytes)
                    response = client.document_text_detection(image=image)
                    page_text = response.full_text_annotation.text if response.full_text_annotation else ""
                    
                    if page_text.strip():
                        full_text.append(page_text)
                
                full_text = '\n\n'.join(full_text)
            
            # Split text into paragraphs
            paragraphs = [p for p in full_text.split('\n') if p.strip()]
            
            # Create chunks of text within token limit
            current_chunk = []
            current_chunk_size = 0
            total_paragraphs = len(paragraphs)
            
            for i, paragraph in enumerate(paragraphs):
                # Update progress (50-90%)
                progress = 50 + int((i / total_paragraphs) * 40)
                progress_bar.progress(progress)
                status_message.write("L'IA résume le texte...")
                
                paragraph_size = len(paragraph)
                
                if current_chunk_size + paragraph_size > MAX_CHUNK_SIZE:
                    # Process current chunk
                    if current_chunk:
                        chunk_text = '\n\n'.join(current_chunk)
                        summary = process_text_with_gpt(
                            prompt=prompt_template_general.format(chunk_text)
                        )
                        if summary:
                            all_chunks_summaries.append(summary)
                    
                    # Start new chunk
                    current_chunk = [paragraph]
                    current_chunk_size = paragraph_size
                else:
                    current_chunk.append(paragraph)
                    current_chunk_size += paragraph_size
            
            # Process final chunk if it exists
            if current_chunk:
                chunk_text = '\n\n'.join(current_chunk)
                summary = process_text_with_gpt(
                    prompt=prompt_template_general.format(chunk_text)
                )
                if summary:
                    all_chunks_summaries.append(summary)
            
            # Combine all summaries and add title
            status_message.write("Finalisation du résumé...")
            progress_bar.progress(90)
            time.sleep(2)
            
            # Get document name without extension
            doc_name = os.path.splitext(uploaded_file.name)[0]
            
            # Create final summary with title (fixed f-string)
            chunks_text = '\n\n'.join(all_chunks_summaries)
            final_summary = f"Résumé - {doc_name}\n\n{chunks_text}"
            
            # Clear progress indicators
            progress_bar.progress(100)
            time.sleep(2)
            progress_bar.empty()
            status_message.empty()
            
            return final_summary
            
        except Exception as e:
            st.error(f"Une erreur est survenue lors du traitement du document: {str(e)}")
            return None

    def process_text_with_gpt(prompt):
        """Handle GPT API calls for single document summarization."""
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini-2024-07-18",
                messages=[{
                    "role": "user",
                    "content": prompt
                }],
                temperature=0
            )
            
            return response.choices[0].message.content.strip()
                
        except Exception as e:
            print(f"Error in GPT processing: {e}")
            return None

    def create_summary_word_document(summary_text, document_name):
        """Create a Word document from the summary text."""
        doc = Document()
        
        # Add title
        title = doc.add_paragraph(f"Résumé - {document_name}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Add spacing after title
        doc.add_paragraph()
        
        # Add summary text
        doc.add_paragraph(summary_text)
        
        # Save to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer

    st.title("RÉSUMÉ DE DOCUMENT PDF OU WORD")
    
    st.markdown("\n")
    st.markdown("#####  Uploadez un fichier PDF ou Word et obtenez son résumé en format .txt ou Word.")

    st.markdown("\n")

    st.markdown("####")
    st.markdown("#####  IMPORTANT")
    st.markdown("######  - Un seul document à la fois.")
    st.markdown("######  - Cet outil n'extrait et ne résume que le texte du document.")
    st.markdown("######  - Formats acceptés : PDF ou Word (.doc, .docx)")
    st.markdown("######  - Pour les PDFs : meilleure qualité d'image possible")
    st.markdown("####")
    
    # Initialize session state for summary
    if 'general_summary' not in st.session_state:
        st.session_state.general_summary = None
    
    st.markdown("#### Ajoutez votre document ci-dessous")
    uploaded_file = st.file_uploader(
        label="Document Upload",  # Add descriptive label
        type=["pdf", "doc", "docx"], 
        accept_multiple_files=False,
        label_visibility="collapsed",
        key="document_uploader_summary"
    )
    
    if uploaded_file:
        if st.button("Générer le résumé"):
            with st.spinner("Génération du résumé en cours..."):
                st.session_state.general_summary = create_single_document_summary(uploaded_file)
    
    if st.session_state.general_summary:
        st.markdown("### ")
        st.markdown("### Résumé")
        st.markdown(st.session_state.general_summary)
        st.markdown("### ")
        
        # Get document name without extension
        doc_name = os.path.splitext(uploaded_file.name)[0]
        
        # Create two columns for download buttons
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                label="📥 Télécharger en .txt",
                data=st.session_state.general_summary,
                file_name=f"resume_{doc_name}.txt",
                mime="text/plain",
                key="download_txt_button"
            )
        
        with col2:
            word_buffer = create_summary_word_document(
                summary_text=st.session_state.general_summary,
                document_name=doc_name
            )
            st.download_button(
                label="📥 Télécharger en Word",
                data=word_buffer,
                file_name=f"resume_{doc_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word_button"
            )

